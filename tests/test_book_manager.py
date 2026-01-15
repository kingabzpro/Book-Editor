import os
import shutil
import tempfile
import json
from pathlib import Path
import pytest

from book_rewriter.book_manager import (
    create_book_structure,
    list_books,
    get_active_book,
    set_active_book,
    delete_book,
    validate_book_structure,
    create_book_name_from_docx,
    get_book_source_path,
    get_book_metadata_path,
    get_book_rewrites_path,
    get_book_validation_path,
    get_book_index_path,
    get_book_config_path,
    BOOKS_DIR,
    REGISTRY_FILE,
    save_central_config,
    load_central_config,
)


@pytest.fixture
def temp_books_dir(tmp_path):
    """Create a temporary books directory for testing."""
    original_dir = os.getcwd()
    temp_books = tmp_path / "books"
    temp_books.mkdir()

    # Temporarily change BOOKS_DIR
    import book_rewriter.book_manager as bm

    original_books_dir = bm.BOOKS_DIR
    bm.BOOKS_DIR = str(temp_books)

    # Temporarily change REGISTRY_FILE
    original_registry = bm.REGISTRY_FILE
    temp_registry = tmp_path / "books_registry.json"
    bm.REGISTRY_FILE = str(temp_registry)

    yield temp_books

    # Cleanup
    bm.BOOKS_DIR = original_books_dir
    bm.REGISTRY_FILE = original_registry
    os.chdir(original_dir)


@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample DOCX file for testing."""
    import docx

    doc = docx.Document()
    doc.add_heading("Chapter 1: The Beginning", 0)
    doc.add_paragraph("Once upon a time...")
    doc.add_heading("Chapter 2: The Middle", 0)
    doc.add_paragraph("And then something happened...")

    docx_path = tmp_path / "test_book.docx"
    doc.save(str(docx_path))
    return docx_path


class TestBookNameGeneration:
    """Test book name generation from DOCX files."""

    def test_simple_filename(self, sample_docx):
        """Test book name from simple filename."""
        book_name = create_book_name_from_docx(str(sample_docx))
        assert book_name == "test_book"

    def test_filename_with_spaces(self, tmp_path):
        """Test book name from filename with spaces."""
        import docx

        doc = docx.Document()
        docx_path = tmp_path / "My Great Book.docx"
        doc.save(str(docx_path))

        book_name = create_book_name_from_docx(str(docx_path))
        assert book_name == "my_great_book"

    def test_filename_with_special_chars(self, tmp_path):
        """Test book name sanitization for special characters."""
        import docx

        doc = docx.Document()
        docx_path = tmp_path / "Book: The *Final* Edition?.docx"
        doc.save(str(docx_path))

        book_name = create_book_name_from_docx(str(docx_path))
        assert ":" not in book_name
        assert "*" not in book_name
        assert "?" not in book_name
        assert "book_the_final_edition" == book_name


class TestBookCreation:
    """Test book creation and structure."""

    def test_create_book_structure(self, temp_books_dir, sample_docx):
        """Test creating a book with proper structure."""
        book_name = "test_book_creation"
        book_info = create_book_structure(book_name, str(sample_docx))

        # Check book directory exists
        assert (temp_books_dir / book_name).exists()

        # Check all subdirectories
        assert get_book_source_path(book_name).exists()
        assert get_book_metadata_path(book_name).exists()
        assert get_book_rewrites_path(book_name).exists()
        assert get_book_validation_path(book_name).exists()
        assert get_book_index_path(book_name).exists()

        # Check source file was copied
        source_files = list(get_book_source_path(book_name).glob("*.docx"))
        assert len(source_files) == 1

        # Check config file
        config_path = get_book_config_path(book_name)
        assert config_path.exists()

        with open(config_path, "r") as f:
            config = json.load(f)
        assert config["book_name"] == book_name

    def test_create_book_duplicate(self, temp_books_dir, sample_docx):
        """Test that duplicate book names are prevented."""
        book_name = "test_duplicate"
        create_book_structure(book_name, str(sample_docx))

        with pytest.raises(ValueError, match="already exists"):
            create_book_structure(book_name, str(sample_docx))


class TestBookListing:
    """Test listing and managing books."""

    def test_list_empty_books(self, temp_books_dir):
        """Test listing when no books exist."""
        books = list_books()
        assert books == []

    def test_list_single_book(self, temp_books_dir, sample_docx):
        """Test listing a single book."""
        create_book_structure("book1", str(sample_docx))
        books = list_books()

        assert len(books) == 1
        assert books[0]["name"] == "book1"
        assert books[0]["is_active"] == True

    def test_list_multiple_books(self, temp_books_dir, sample_docx):
        """Test listing multiple books."""
        create_book_structure("book1", str(sample_docx))
        create_book_structure("book2", str(sample_docx))
        create_book_structure("book3", str(sample_docx))

        books = list_books()
        assert len(books) == 3
        book_names = [b["name"] for b in books]
        assert "book1" in book_names
        assert "book2" in book_names
        assert "book3" in book_names

    def test_active_book_tracking(self, temp_books_dir, sample_docx):
        """Test active book tracking."""
        create_book_structure("book1", str(sample_docx))
        assert get_active_book() == "book1"

        create_book_structure("book2", str(sample_docx))
        assert get_active_book() == "book1"  # First book stays active


class TestBookDeletion:
    """Test book deletion."""

    def test_delete_book(self, temp_books_dir, sample_docx):
        """Test deleting a book."""
        book_name = "test_delete"
        create_book_structure(book_name, str(sample_docx))

        # Verify book exists
        assert (temp_books_dir / book_name).exists()
        assert len(list_books()) == 1

        # Delete book
        delete_book(book_name)

        # Verify book is gone
        assert not (temp_books_dir / book_name).exists()
        assert len(list_books()) == 0

    def test_delete_nonexistent_book(self, temp_books_dir):
        """Test deleting a book that doesn't exist."""
        with pytest.raises(ValueError, match="not found"):
            delete_book("nonexistent_book")


class TestBookValidation:
    """Test book structure validation."""

    def test_validate_complete_book(self, temp_books_dir, sample_docx):
        """Test validation of complete book structure."""
        book_name = "test_valid"
        create_book_structure(book_name, str(sample_docx))

        result = validate_book_structure(book_name)
        assert result["valid"] == True
        assert len(result["issues"]) == 0

    def test_validate_missing_dirs(self, temp_books_dir, sample_docx):
        """Test validation detects missing directories."""
        book_name = "test_invalid"
        create_book_structure(book_name, str(sample_docx))

        # Remove a directory
        shutil.rmtree(get_book_source_path(book_name))

        result = validate_book_structure(book_name)
        assert result["valid"] == False
        assert any("Missing directory" in issue for issue in result["issues"])

    def test_validate_nonexistent_book(self, temp_books_dir):
        """Test validation of nonexistent book."""
        result = validate_book_structure("nonexistent")
        assert result["valid"] == False
        assert "not found" in result["issues"][0]


class TestCentralConfig:
    """Test central configuration management."""

    def test_save_and_load_config(self, tmp_path):
        """Test saving and loading central config."""
        import book_rewriter.book_manager as bm

        original_config_file = bm.CENTRAL_CONFIG_FILE
        temp_config_file = tmp_path / "central_config.json"
        bm.CENTRAL_CONFIG_FILE = str(temp_config_file)

        config = {
            "api_keys": {
                "nebius": "test_key",
                "mistral": "test_key2",
            },
            "default_settings": {
                "target_word_count_min": 1500,
            },
        }

        save_central_config(config)
        loaded = load_central_config()

        assert loaded == config

        # Cleanup
        bm.CENTRAL_CONFIG_FILE = original_config_file

    def test_load_missing_config(self, tmp_path):
        """Test loading config when file doesn't exist."""
        import book_rewriter.book_manager as bm

        original_config_file = bm.CENTRAL_CONFIG_FILE
        temp_config_file = tmp_path / "nonexistent_config.json"
        bm.CENTRAL_CONFIG_FILE = str(temp_config_file)

        loaded = load_central_config()
        assert loaded == {}

        # Cleanup
        bm.CENTRAL_CONFIG_FILE = original_config_file


class TestPathResolution:
    """Test path resolution functions."""

    def test_get_book_paths(self, temp_books_dir, sample_docx):
        """Test getting various book paths."""
        book_name = "test_paths"
        create_book_structure(book_name, str(sample_docx))

        source_path = get_book_source_path(book_name)
        metadata_path = get_book_metadata_path(book_name)
        rewrites_path = get_book_rewrites_path(book_name)
        validation_path = get_book_validation_path(book_name)
        index_path = get_book_index_path(book_name)
        config_path = get_book_config_path(book_name)

        assert book_name in str(source_path)
        assert book_name in str(metadata_path)
        assert book_name in str(rewrites_path)
        assert book_name in str(validation_path)
        assert book_name in str(index_path)
        assert book_name in str(config_path)

        assert source_path.name == "source"
        assert metadata_path.name == "metadata"
        assert rewrites_path.name == "rewrites"
        assert validation_path.name == "validation"
        assert index_path.name == "index"


def test_sanitize_name():
    """Test name sanitization."""
    from book_rewriter.book_manager import sanitize_name

    assert sanitize_name("My Book") == "my_book"
    assert sanitize_name("Book: The Final") == "book_the_final"
    # After sanitization, * and ? are removed, so Test*File?.docx becomes testfile.docx
    assert sanitize_name("Test*File?.docx") == "testfile.docx"
    assert sanitize_name("A" * 100) == "a" * 50  # Length limit


def test_ensure_books_dir(tmp_path):
    """Test ensuring books directory exists."""
    from book_rewriter.book_manager import ensure_books_dir
    import book_rewriter.book_manager as bm

    original_books_dir = bm.BOOKS_DIR
    bm.BOOKS_DIR = str(tmp_path / "test_books")

    ensure_books_dir()
    assert Path(bm.BOOKS_DIR).exists()

    # Cleanup
    bm.BOOKS_DIR = original_books_dir
